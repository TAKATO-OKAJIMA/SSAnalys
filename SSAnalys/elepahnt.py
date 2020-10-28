import requests
import sys
from bs4 import BeautifulSoup
from docx import Document

if __name__ == '__main__':
    url = sys.argv[1]
    html = requests.get(url)
    soup = BeautifulSoup(html.content, 'html.parser')

    title = soup.select_one('title')
    contents = soup.select('#main > div.article > dl > dd')
    shortstory = []

    for content in contents:
        shortstory.append(content.text)

    print(shortstory)
    document = Document()

    for storypart in shortstory:
        storylines = storypart.split('   ')
        for storyline in storylines:
            document.add_paragraph(storyline).add_run('\n')

    document.save(title.text + '.docx')